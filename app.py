from flask import Flask, request, send_file, render_template
from reportlab.platypus import SimpleDocTemplate, Paragraph, Image, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
import base64
import uuid
import os
from PIL import Image as PILImage, ImageDraw
from reportlab.platypus import SimpleDocTemplate, Paragraph, Image, Spacer, HRFlowable
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from datetime import datetime
from flask import request, send_file
from reportlab.pdfgen import canvas




app = Flask(__name__)



def process_hero_image(input_path, output_path):
    img = PILImage.open(input_path).convert("RGB")

    width, height = img.size

    # ---- Crop to banner ratio (16:6) ----
    target_ratio = 16 / 6
    new_height = int(width / target_ratio)

    if new_height < height:
        top = (height - new_height) // 2
        bottom = top + new_height
        img = img.crop((0, top, width, bottom))
    else:
        new_width = int(height * target_ratio)
        left = (width - new_width) // 2
        right = left + new_width
        img = img.crop((left, 0, right, height))

    # ---- Resize to consistent size ----
    img = img.resize((1000, 375))

    # ---- Rounded corners ----
    radius = 30
    mask = PILImage.new("L", img.size, 0)
    draw = ImageDraw.Draw(mask)

    draw.rounded_rectangle(
        [(0, 0), img.size],
        radius=radius,
        fill=255
    )

    img.putalpha(mask)

    # ---- Add white background (for PDF compatibility) ----
    bg = PILImage.new("RGB", img.size, (255, 255, 255))
    bg.paste(img, mask=img.split()[3])

    bg.save(output_path)

    return output_path



# create temp folder for images
if not os.path.exists("temp"):
    os.makedirs("temp")
def process_hero_image(input_path, output_path):
    img = PILImage.open(input_path).convert("RGB")
    width, height = img.size
    # ---- Crop to banner ratio (16:6) ----
    target_ratio = 16 / 6
    new_height = int(width / target_ratio)
    if new_height < height:
        top = (height - new_height) // 2
        bottom = top + new_height
        img = img.crop((0, top, width, bottom))
    else:
        new_width = int(height * target_ratio)
        left = (width - new_width) // 2
        right = left + new_width
        img = img.crop((left, 0, right, height))

    # ---- Resize (consistent output) ----
    img = img.resize((1000, 375))

    # ---- Add rounded corners ----
    radius = 30
    mask = PILImage.new("L", img.size, 0)
    draw = ImageDraw.Draw(mask)
    draw.rounded_rectangle(
        [(0, 0), img.size],
        radius=radius,
        fill=255
    )

    img.putalpha(mask)

    # white background to support PDF
    bg = PILImage.new("RGB", img.size, (255, 255, 255))
    bg.paste(img, mask=img.split()[3])

    bg.save(output_path)

    return output_path


# convert base64 image → file
def save_base64_image(data):
    header, encoded = data.split(",", 1)
    filename = f"temp/{uuid.uuid4()}.jpg"
    with open(filename, "wb") as f:
        f.write(base64.b64decode(encoded))
    return filename


# home page
@app.route('/')
def home():
    return render_template('index.html')


# ---- Background Canvas ----
# class DarkCanvas(canvas.Canvas):
#     def set_background(self):
#         self.setFillColorRGB(0.06, 0.09, 0.16)  # dark blue
#         self.rect(0, 0, self._pagesize[0], self._pagesize[1], fill=1)

    def showPage(self):
        self.set_background()
        canvas.Canvas.showPage(self)

    def save(self):
        self.set_background()
        canvas.Canvas.save(self)


@app.route('/export-pdf', methods=['POST'])
def export_pdf():

    doc = SimpleDocTemplate(
        "blog.pdf",
        rightMargin=50,
        leftMargin=50,
        topMargin=40,
        bottomMargin=40
    )

    elements = []

    # ---- STYLES ----
    category_style = ParagraphStyle(
        name='Category',
        fontSize=10,
        textColor=colors.HexColor("#2563eb"),  # blue
        spaceAfter=8
    )

    title_style = ParagraphStyle(
        name='Title',
        fontSize=28,
        leading=32,
        textColor=colors.HexColor("#f97316"),  # near black
        spaceAfter=15
    )

    meta_style = ParagraphStyle(
        name='Meta',
        fontSize=10,
        textColor=colors.HexColor("#6b7280"),  # gray
        spaceAfter=15
    )

    desc_style = ParagraphStyle(
        name='Desc',
        fontSize=12,
        leading=18,
        textColor=colors.HexColor("#4b5563"),  # soft gray
        spaceAfter=25
    )

    heading_style = ParagraphStyle(
        name='Heading',
        fontSize=18,
        leading=22,
        textColor=colors.HexColor("#1f2937"),  # dark gray
        spaceAfter=10
    )

    content_style = ParagraphStyle(
        name='Content',
        fontSize=12,
        leading=18,
        textColor=colors.HexColor("#111827"),  # black
        spaceAfter=15
    )

    # ---- DATA ----
    title = request.form.get("title")
    desc = request.form.get("description")
    hero = request.files.get("hero")

    os.makedirs("temp", exist_ok=True)

    # ---- HERO IMAGE ----
    if hero and hero.filename:
        original_path = f"temp/original_{hero.filename}"
        processed_path = f"temp/processed_{hero.filename}"

        hero.save(original_path)

        process_hero_image(original_path, processed_path)

        img = Image(processed_path, width=500, height=190)
        img.hAlign = 'CENTER'

        elements.append(img)
        elements.append(Spacer(1, 20))

    # ---- CATEGORY ----
    # elements.append(Paragraph("INSIGHT", category_style))

    # ---- TITLE ----
    elements.append(Paragraph(title, title_style))

    # ---- META ----
    # elements.append(Paragraph("Published on March 28, 2026", meta_style))

    # ---- DESCRIPTION ----
    elements.append(Paragraph(desc, desc_style))

    # ---- SECTIONS ----
    i = 0
    while True:
        heading = request.form.get(f"heading_{i}")
        content = request.form.get(f"content_{i}")
        image = request.files.get(f"image_{i}")

        if not heading and not content:
            break

        elements.append(Paragraph(heading, heading_style))

        if image:
            path = f"temp/{image.filename}"
            image.save(path)

            img = Image(path, width=420, height=240)
            img.hAlign = 'CENTER'
            elements.append(img)
            elements.append(Spacer(1, 15))

        elements.append(Paragraph(content, content_style))
        elements.append(Spacer(1, 20))

        i += 1

    doc.build(elements)

    return send_file("blog.pdf", as_attachment=True)



@app.route('/export-docx', methods=['POST'])
def export_docx():
    from docx import Document

    doc = Document()

    title = request.form.get("title")
    desc = request.form.get("description")
    hero = request.files.get("hero")

    doc.add_heading(title, 0)
    doc.add_paragraph(desc)

    if hero and hero.filename:
        path = f"temp/{hero.filename}"
        hero.save(path)

        doc.add_picture(path)

    i = 0
    while True:
        heading = request.form.get(f"heading_{i}")
        content = request.form.get(f"content_{i}")
        image = request.files.get(f"image_{i}")

        if not heading and not content:
            break

        doc.add_heading(heading, level=1)
        doc.add_paragraph(content)

        if image:
            path = f"temp/{image.filename}"
            image.save(path)
            doc.add_picture(path)

        i += 1

    doc.save("blog.docx")
    return send_file("blog.docx", as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)